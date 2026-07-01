import streamlit as st
from anthropic import Anthropic
from dotenv import load_dotenv
import docx
from docx import Document
from docx.shared import Pt, Inches
import io
import difflib
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
import textwrap
import json
import os
from datetime import date

load_dotenv()
client = Anthropic()

TRACKER_FILE = "job_tracker.json"

st.set_page_config(
    page_title="AI Job Application Assistant",
    page_icon="💼",
    layout="wide"
)

st.markdown("""
    <style>
        .stApp {
            background-color: #f0f4ff;
            color: #1a1a2e;
        }
        [data-testid="stSidebar"] {
            background-color: #1a3a6b;
            border-right: 1px solid #1a3a6b;
        }
        [data-testid="stSidebar"] * {
            color: #ffffff !important;
        }
        [data-testid="stSidebar"] hr {
            border-color: #ffffff44;
        }
        .stTabs [data-baseweb="tab-list"] {
            background-color: #dce6ff;
            border-radius: 8px;
            padding: 4px;
        }
        .stTabs [data-baseweb="tab"] {
            color: #1a3a6b;
            font-weight: 500;
        }
        .stTabs [data-baseweb="tab"][aria-selected="true"] {
            background-color: #7aa4f0 !important;
            color: #ffffff !important;
            border-radius: 6px !important;
            font-weight: 700 !important;
        }
        .stTabs [data-baseweb="tab-highlight"] {
            background-color: #7aa4f0 !important;
        }
        .stTabs [data-baseweb="tab-border"] {
            background-color: #7aa4f0 !important;
        }
        .stButton > button {
            background-color: #4f7de0;
            color: #ffffff !important;
            border: none;
            border-radius: 8px;
            padding: 10px 24px;
            font-weight: 600;
            font-size: 15px;
            width: 100%;
            transition: background-color 0.2s;
        }
        .stButton > button:hover {
            background-color: #14aff0;
            color: #ffffff !important;
        }
        .stTextArea textarea {
            background-color: #ffffff;
            color: #1a1a2e !important;
            border: 1px solid #b0c4f0;
            border-radius: 8px;
            font-size: 14px;
        }
        .stTextInput input {
            background-color: #ffffff;
            color: #1a1a2e !important;
            border: 1px solid #b0c4f0;
            border-radius: 8px;
        }
        .stSelectbox div {
            background-color: #ffffff;
            color: #1a1a2e !important;
            border-radius: 8px;
        }
        [data-testid="stMetric"] {
            background-color: #ffffff;
            border: 1px solid #b0c4f0;
            border-radius: 10px;
            padding: 16px;
        }
        [data-testid="stMetricValue"] {
            color: #1a3a6b !important;
            font-weight: 700;
        }
        [data-testid="stMetricLabel"] {
            color: #4a6a9b !important;
        }
        h1 {
            color: #1a3a6b !important;
            font-size: 28px;
            font-weight: 700;
        }
        h2, h3 {
            color: #1a3a6b !important;
            font-weight: 600;
        }
        p, span, label, div {
            color: #1a1a2e;
        }
        .streamlit-expanderHeader {
            background-color: #ffffff;
            border: 1px solid #b0c4f0;
            border-radius: 8px;
            color: #1a3a6b !important;
        }
        .stDownloadButton > button {
            background-color: #0a8f4f;
            color: #ffffff !important;
            border-radius: 8px;
            font-weight: 600;
            width: 100%;
        }
        .stDownloadButton > button:hover {
            background-color: #077a42;
            color: #ffffff !important;
        }
        [data-testid="stFileUploader"] {
            background-color: #ffffff;
            border: 1px dashed #b0c4f0;
            border-radius: 8px;
            padding: 8px;
        }
        hr {
            border-color: #b0c4f0;
        }
        .stRadio label {
            color: #ffffff !important;
        }
        .stSuccess {
            background-color: #e6f9ee;
            border: 1px solid #0a8f4f;
            color: #1a1a2e !important;
        }
        .stWarning {
            background-color: #fff8e6;
            border: 1px solid #f0a500;
            color: #1a1a2e !important;
        }
        .stInfo {
            background-color: #e6f0ff;
            border: 1px solid #1a3a6b;
            color: #1a1a2e !important;
        }
    </style>
""", unsafe_allow_html=True)

def load_tracker():
    if os.path.exists(TRACKER_FILE):
        with open(TRACKER_FILE, "r") as f:
            return json.load(f)
    return []

def save_tracker(data):
    with open(TRACKER_FILE, "w") as f:
        json.dump(data, f, indent=2)

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_from_pdf(file):
    import pypdf
    reader = pypdf.PdfReader(file)
    return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def create_docx(text):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    for line in text.split("\n"):
        line = line.strip()
        if line == "":
            doc.add_paragraph("")
        elif line.isupper() or (line.endswith(":") and len(line) < 50):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("-") or line.startswith("•"):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(line.lstrip("-•").strip())
            run.font.size = Pt(11)
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(11)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def create_pdf(text):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 1 * inch
    y = height - margin
    c.setFont("Helvetica", 11)
    for line in text.split("\n"):
        if y < margin:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin
        if line.strip() == "":
            y -= 6
        else:
            wrapped = textwrap.wrap(line, width=90)
            if not wrapped:
                y -= 6
            for wrap_line in wrapped:
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - margin
                c.drawString(margin, y, wrap_line)
                y -= 16
    c.save()
    buffer.seek(0)
    return buffer.read()

def get_prompt(mode, job_description, resume_text):
    if mode == "Full Resume":
        return (
            "Rewrite this resume to better match the job description. "
            "Tailor the bullet points and skills section. "
            "Keep the exact same structure, sections, and formatting. "
            "Keep it professional and realistic.\n\n"
            f"Job Description:\n{job_description}\n\n"
            f"Resume:\n{resume_text}\n\n"
            "Return only the full rewritten resume, nothing else. "
            "Use dashes (-) for all bullet points."
        )
    else:
        return (
            "Rewrite ONLY the bullet points and skills section of this resume to better match the job description. "
            "Do not change the name, contact info, job titles, company names, or dates. "
            "Keep everything else exactly the same. "
            "Keep it professional and realistic.\n\n"
            f"Job Description:\n{job_description}\n\n"
            f"Resume:\n{resume_text}\n\n"
            "Return the full resume with only the bullet points and skills section rewritten. "
            "Use dashes (-) for all bullet points."
        )

def get_match_score(job_description, resume_text):
    message = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=500,
        messages=[
            {
                "role": "user",
                "content": (
                    "Compare this resume to the job description and return a JSON object with exactly these fields:\n"
                    "- score: a number from 0 to 100 representing how well the resume matches the job\n"
                    "- missing: a list of 3 to 5 short keywords or skills that are in the job description but missing from the resume\n"
                    "Return only the JSON object, nothing else.\n\n"
                    f"Job Description:\n{job_description}\n\n"
                    f"Resume:\n{resume_text}"
                )
            }
        ]
    )
    raw = message.content[0].text.strip()
    clean = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(clean)

def show_diff(original, rewritten):
    original_lines = original.splitlines()
    rewritten_lines = rewritten.splitlines()
    diff = difflib.ndiff(original_lines, rewritten_lines)
    html = "<div style='font-family: monospace; font-size: 13px; line-height: 1.6;'>"
    for line in diff:
        if line.startswith("- "):
            html += f"<div style='background-color:#ffe6e6; color:#cc0000; padding:2px 8px; border-radius:4px; margin:1px 0;'>❌ {line[2:]}</div>"
        elif line.startswith("+ "):
            html += f"<div style='background-color:#e6ffee; color:#007a33; padding:2px 8px; border-radius:4px; margin:1px 0;'>✅ {line[2:]}</div>"
        elif line.startswith("  "):
            html += f"<div style='padding:2px 8px; color:#1a1a2e;'>{line[2:]}</div>"
    html += "</div>"
    return html

if "result" not in st.session_state:
    st.session_state.result = None
if "before_score" not in st.session_state:
    st.session_state.before_score = None
if "after_score" not in st.session_state:
    st.session_state.after_score = None
if "resume_text" not in st.session_state:
    st.session_state.resume_text = None
if "rewrite_mode" not in st.session_state:
    st.session_state.rewrite_mode = None

with st.sidebar:
    st.markdown("## ⚙️ Options")
    st.markdown("---")
    rewrite_mode = st.radio(
        "Rewrite Mode",
        ["Full Resume", "Bullets & Skills Only"],
        help="Full Resume rewrites everything. Bullets & Skills Only keeps your structure intact."
    )
    st.markdown("---")
    output_format = st.radio(
        "Output Format",
        ["Word Doc (.docx)", "PDF"],
        help="Word Doc is easier to edit. PDF is ready to send."
    )
    st.markdown("---")
    st.markdown("### 💡 Tips")
    st.markdown("- Upload a **.docx** for best results")
    st.markdown("- Use **Bullets & Skills Only** to preserve formatting")
    st.markdown("- Check **What Changed** to review edits")
    st.markdown("- Save to **Job Tracker** after each application")

st.markdown("# 💼 AI Job Application Assistant")
st.markdown("*Tailor your resume to any job in seconds.*")
st.markdown("---")

tab1, tab2 = st.tabs(["📄 Resume Rewriter", "📊 Job Tracker"])

with tab1:
    col_left, col_right = st.columns(2)

    with col_left:
        st.markdown("### 📋 Job Description")
        job_description = st.text_area(
            "Job Description",
            height=250,
            placeholder="Copy and paste the full job posting here...",
            label_visibility="collapsed"
        )

    with col_right:
        st.markdown("### 📁 Your Resume")
        uploaded_file = st.file_uploader(
            "Upload your resume",
            type=["pdf", "docx"],
            label_visibility="collapsed"
        )
        pasted_resume = st.text_area(
            "Paste your resume",
            height=180,
            placeholder="Paste your resume here if you don't have a file...",
            label_visibility="collapsed"
        )

    st.markdown("---")

    if st.button("✨ Rewrite My Resume"):
        resume_text = ""

        if uploaded_file is not None:
            if uploaded_file.name.endswith(".docx"):
                resume_text = extract_text_from_docx(uploaded_file)
            elif uploaded_file.name.endswith(".pdf"):
                resume_text = extract_text_from_pdf(uploaded_file)
        elif pasted_resume.strip():
            resume_text = pasted_resume.strip()

        if not job_description.strip():
            st.warning("⚠️ Please enter a job description.")
        elif not resume_text:
            st.warning("⚠️ Please upload a resume file or paste your resume text.")
        else:
            before_score = get_match_score(job_description, resume_text)

            with st.spinner("✨ Analyzing and rewriting your resume..."):
                message = client.messages.create(
                    model="claude-haiku-4-5",
                    max_tokens=2000,
                    messages=[
                        {
                            "role": "user",
                            "content": get_prompt(rewrite_mode, job_description, resume_text)
                        }
                    ]
                )
                result = message.content[0].text

            after_score = get_match_score(job_description, result)

            st.session_state.result = result
            st.session_state.before_score = before_score
            st.session_state.after_score = after_score
            st.session_state.resume_text = resume_text
            st.session_state.rewrite_mode = rewrite_mode

    if st.session_state.result:
        result = st.session_state.result
        before_score = st.session_state.before_score
        after_score = st.session_state.after_score
        resume_text = st.session_state.resume_text
        rewrite_mode = st.session_state.rewrite_mode

        st.markdown("---")
        st.markdown("### 📊 Match Score")
        col1, col2 = st.columns(2)
        with col1:
            st.metric(label="Before Rewrite", value=f"{before_score['score']}%")
        with col2:
            st.metric(
                label="After Rewrite",
                value=f"{after_score['score']}%",
                delta=f"+{after_score['score'] - before_score['score']}%"
            )

        if after_score["missing"]:
            st.markdown("### 🔍 Still Missing From Your Resume")
            cols = st.columns(len(after_score["missing"]))
            for i, item in enumerate(after_score["missing"]):
                with cols[i]:
                    st.markdown(f"<div style='background-color:#e6f0ff; border:1px solid #b0c4f0; border-radius:8px; padding:8px 12px; text-align:center; color:#1a3a6b;'>⚠️ {item}</div>", unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 📝 Your Rewritten Resume")
        st.write(result)

        st.markdown("---")
        with st.expander("🔍 See What Changed"):
            diff_html = show_diff(resume_text, result)
            st.markdown(diff_html, unsafe_allow_html=True)

        st.markdown("---")
        if output_format == "Word Doc (.docx)":
            file_data = create_docx(result)
            st.download_button(
                label="⬇️ Download as Word Doc",
                data=file_data,
                file_name="rewritten_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            file_data = create_pdf(result)
            st.download_button(
                label="⬇️ Download as PDF",
                data=file_data,
                file_name="rewritten_resume.pdf",
                mime="application/pdf"
            )

        st.markdown("---")
        st.markdown("### 💾 Save to Job Tracker")
        col_a, col_b = st.columns(2)
        with col_a:
            job_title = st.text_input("Job Title", placeholder="e.g. Systems Administrator", key="job_title_input")
        with col_b:
            company = st.text_input("Company Name", placeholder="e.g. Microsoft", key="company_input")
        status = st.selectbox("Application Status", ["Applied", "Interview", "Offer", "Rejected"])

        if st.button("💾 Save to Tracker"):
            if job_title and company:
                tracker_data = load_tracker()
                tracker_data.append({
                    "date": str(date.today()),
                    "job_title": job_title,
                    "company": company,
                    "status": status,
                    "rewrite_mode": rewrite_mode,
                    "before_score": before_score["score"],
                    "after_score": after_score["score"]
                })
                save_tracker(tracker_data)
                st.success("✅ Saved to Job Tracker!")
            else:
                st.warning("⚠️ Please enter a job title and company name.")

with tab2:
    st.markdown("### 📊 Your Job Applications")
    tracker = load_tracker()

    if not tracker:
        st.info("📭 No applications tracked yet. Rewrite a resume and save it to start tracking.")
    else:
        st.markdown(f"**{len(tracker)} application(s) tracked**")
        st.markdown("---")
        for i, job in enumerate(reversed(tracker)):
            status_colors = {
                "Applied": "#4f7de0",
                "Interview": "#f0a500",
                "Offer": "#0a8f4f",
                "Rejected": "#e05252"
            }
            color = status_colors.get(job["status"], "#4f7de0")
            with st.expander(f"**{job['job_title']}** at {job['company']} — {job['date']}"):
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.markdown(f"<div style='background-color:{color}22; border:1px solid {color}; border-radius:8px; padding:8px 12px; text-align:center; color:{color}; font-weight:600;'>{job['status']}</div>", unsafe_allow_html=True)
                with col2:
                    st.metric(label="Before Score", value=f"{job['before_score']}%")
                with col3:
                    st.metric(label="After Score", value=f"{job['after_score']}%")

                st.write(f"**Rewrite Mode:** {job['rewrite_mode']}")
                st.markdown("---")
                new_status = st.selectbox(
                    "Update Status",
                    ["Applied", "Interview", "Offer", "Rejected"],
                    index=["Applied", "Interview", "Offer", "Rejected"].index(job["status"]),
                    key=f"status_{i}"
                )
                if st.button("Update Status", key=f"update_{i}"):
                    actual_index = len(tracker) - 1 - i
                    tracker[actual_index]["status"] = new_status
                    save_tracker(tracker)
                    st.success("✅ Status updated!")
                    st.rerun()